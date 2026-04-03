#!/usr/bin/env python3
"""
generate_agreement.py — Symphony Smart Homes Agreement Addendum Generator

Generates a formatted .docx addendum from CLI arguments.

Usage:
    python generate_agreement.py \\
        --client "John Smith" \\
        --project "123 Main St Smart Home" \\
        --items "65-inch Samsung QLED TV, Sonos Arc soundbar" \\
        --integrations "Control4 OS 3.x, Lutron RadioRA 3, Sonos S2" \\
        --support-days 90

Outputs the path of the generated .docx file to stdout.
"""

import argparse
import os
import sys
from datetime import date

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("ERROR: python-docx is not installed. Run: pip install python-docx", file=sys.stderr)
    sys.exit(1)


# ── Styling helpers ────────────────────────────────────────────────────────────

BRAND_DARK = RGBColor(0x1C, 0x2B, 0x3A)   # dark navy
BRAND_TEAL = RGBColor(0x01, 0x69, 0x6F)   # Symphony teal
MUTED_GRAY = RGBColor(0x7A, 0x79, 0x74)


def set_font(run, bold=False, size=11, color=None, italic=False):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    """Add a styled section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(13 if level == 1 else 11)
    run.bold = True
    run.font.color.rgb = BRAND_TEAL
    return p


def add_body(doc, text, space_after=6):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(space_after)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)
        run.font.color.rgb = BRAND_DARK
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    run.font.color.rgb = BRAND_DARK
    p.paragraph_format.space_after = Pt(2)
    return p


def add_label_value(doc, label, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r_label = p.add_run(f"{label}: ")
    r_label.bold = True
    r_label.font.name = "Calibri"
    r_label.font.size = Pt(10.5)
    r_label.font.color.rgb = BRAND_DARK
    r_val = p.add_run(value)
    r_val.font.name = "Calibri"
    r_val.font.size = Pt(10.5)
    r_val.font.color.rgb = BRAND_DARK
    return p


def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "D4D1CA")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def add_signature_block(doc, party_label):
    add_heading(doc, party_label, level=2)
    for field in ["Name (Print)", "Signature", "Title", "Date"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(14)
        r = p.add_run(f"{field}: ")
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(10.5)
        r.font.color.rgb = BRAND_DARK
        r_line = p.add_run("_" * 45)
        r_line.font.name = "Calibri"
        r_line.font.size = Pt(10.5)
        r_line.font.color.rgb = MUTED_GRAY


# ── Document builder ───────────────────────────────────────────────────────────

def build_document(client, project, items, support_days, integrations):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ── Cover header ──────────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(4)
    tr = title_p.add_run("AGREEMENT ADDENDUM")
    tr.font.name = "Calibri"
    tr.font.size = Pt(18)
    tr.bold = True
    tr.font.color.rgb = BRAND_DARK

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.paragraph_format.space_after = Pt(14)
    sr = sub_p.add_run("Symphony Smart Homes")
    sr.font.name = "Calibri"
    sr.font.size = Pt(13)
    sr.font.color.rgb = BRAND_TEAL

    add_horizontal_rule(doc)

    # ── Meta block ────────────────────────────────────────────────────────────
    today_str = date.today().strftime("%B %d, %Y")
    add_label_value(doc, "Client", client)
    add_label_value(doc, "Project", project)
    add_label_value(doc, "Addendum Date", today_str)
    add_label_value(doc, "Addendum Reference", "ADDENDUM-001")

    add_body(
        doc,
        "This Addendum supplements and is incorporated into the Master Services Agreement or "
        "Proposal between Symphony Smart Homes and the Client identified above. In the event "
        "of a conflict between this Addendum and any prior agreement, this Addendum controls "
        "with respect to the subject matter herein.",
        space_after=4,
    )
    add_horizontal_rule(doc)

    # ── Section 1: Client-Supplied Equipment ──────────────────────────────────
    add_heading(doc, "Section 1: Client-Supplied Equipment")
    add_body(doc, "Client is providing the following equipment independently of this agreement:")

    # Items block
    items_list = [i.strip() for i in items.split(",") if i.strip()]
    for item in items_list:
        add_bullet(doc, item)

    add_body(
        doc,
        "Symphony Smart Homes assumes no responsibility for the functionality, compatibility, "
        "warranty, or performance of client-supplied equipment. Installation and integration "
        "of client-supplied equipment is performed on a best-effort basis.",
    )
    add_body(
        doc,
        "If client-supplied equipment is found to be incompatible, defective, or otherwise "
        "insufficient for the intended integration, Symphony Smart Homes will notify the Client "
        "in writing and may require replacement or upgrade of the affected equipment at the "
        "Client's sole expense before proceeding with installation.",
    )
    add_body(
        doc,
        "Any additional labor required as a result of issues arising from client-supplied "
        "equipment will be billed at Symphony Smart Homes' standard service rates in effect "
        "at the time services are rendered.",
    )
    add_horizontal_rule(doc)

    # ── Section 2: TV Installation Services ───────────────────────────────────
    add_heading(doc, "Section 2: TV Installation Services")
    add_body(
        doc,
        "TV installation, mounting, cable management, and Control4 integration constitute a "
        "separate service engagement billed at standard labor rates and are not included in "
        "the base agreement unless explicitly stated in the associated proposal or change order.",
    )

    add_heading(doc, "Included services:", level=2)
    for item in [
        "Physical mounting of television(s) to wall or mount structure",
        "Concealed or routed cable management (HDMI, power, IR/RS-232 control lines)",
        "Control4 driver configuration and source switching integration",
        "Basic picture mode calibration (brightness, contrast, input labeling)",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Excluded services:", level=2)
    for item in [
        "Television procurement, delivery, and unboxing (Client responsibility)",
        "Manufacturer warranty registration (Client responsibility)",
        "Content platform setup, app configuration, or streaming account login",
        "Acoustic calibration or advanced ISF/Dolby Vision calibration (available separately)",
    ]:
        add_bullet(doc, item)

    add_body(
        doc,
        "Detailed pricing for TV installation will be provided in a separate change order "
        "upon confirmation of TV model(s), quantities, and mounting locations.",
    )
    add_horizontal_rule(doc)

    # ── Section 3: Post-Installation Support ──────────────────────────────────
    add_heading(doc, "Section 3: Post-Installation Support")
    add_body(
        doc,
        f"Symphony Smart Homes provides a {support_days}-day post-installation defect coverage "
        "period commencing on the date of system commissioning, as documented in the project "
        "completion sign-off.",
    )

    add_heading(doc, "Coverage includes:", level=2)
    for item in [
        "Incorrect wiring terminations identified after commissioning",
        "Programming errors or logic defects in Control4 or associated automation platforms",
        "Configuration issues that prevent the system from operating as specified",
        "Component failures attributable to the installation process",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Coverage excludes:", level=2)
    for item in [
        "Client-supplied equipment (see Section 1)",
        "Issues arising from third-party firmware updates, cloud service changes, or manufacturer API modifications occurring after commissioning",
        "Physical damage to equipment or infrastructure occurring after installation completion",
        "Modifications made by the Client or any third party to the installed system",
    ]:
        add_bullet(doc, item)

    add_body(
        doc,
        "Support requests outside the scope of defect coverage, or submitted after the coverage "
        "period expires, will be billed at Symphony Smart Homes' standard service rates.",
    )
    add_horizontal_rule(doc)

    # ── Section 4: Third-Party Integration Disclaimers ────────────────────────
    add_heading(doc, "Section 4: Third-Party Integration Disclaimers")
    add_body(
        doc,
        "This system integrates one or more third-party products and services whose "
        "functionality, reliability, and API availability are governed solely by their "
        "respective manufacturers and service providers. Symphony Smart Homes does not "
        "warrant, guarantee, or control the ongoing behavior of third-party platforms.",
    )

    add_heading(doc, "Symphony Smart Homes expressly does not warrant:", level=2)
    for item in [
        "Continued compatibility following manufacturer-issued firmware or software updates",
        "Uninterrupted availability of cloud services, mobile apps, or voice platform integrations",
        "Compatibility following manufacturer changes to APIs, authentication protocols, or device drivers",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Third-party integrations included in this project:", level=2)
    integrations_list = [i.strip() for i in integrations.split(",") if i.strip()]
    for integration in integrations_list:
        add_bullet(doc, integration)

    add_body(
        doc,
        "In the event that a third-party update degrades or breaks an integration, Symphony "
        "Smart Homes will provide restoration assistance at its standard service rates. Clients "
        "may optionally enroll in a Managed Services Agreement for proactive monitoring and "
        "update management.",
    )
    add_horizontal_rule(doc)

    # ── Section 5: Warranty Scope ─────────────────────────────────────────────
    add_heading(doc, "Section 5: Warranty Scope")
    add_body(
        doc,
        "Symphony Smart Homes provides a one (1) year workmanship warranty commencing on the "
        "date of system commissioning.",
    )

    add_heading(doc, "Warranty covers:", level=2)
    for item in [
        "Wire terminations and low-voltage connections installed by Symphony Smart Homes",
        "Cable pathways, conduit, and physical infrastructure installed by Symphony Smart Homes",
        "Equipment mounting and rack installation performed by Symphony Smart Homes",
        "Control4 programming and automation logic delivered as part of this project scope",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Warranty does not cover:", level=2)
    for item in [
        "Equipment hardware failures (subject to individual manufacturer warranties; Symphony will assist with claims coordination)",
        "Client-supplied equipment (see Section 1)",
        "Normal wear and tear on consumable components",
        "Damage resulting from power surges, electrical faults, or acts of God",
        "Modifications, alterations, or additions made by any party other than Symphony Smart Homes",
    ]:
        add_bullet(doc, item)

    add_horizontal_rule(doc)

    # ── Signature block ───────────────────────────────────────────────────────
    add_heading(doc, "Acknowledgment & Signatures")
    add_body(
        doc,
        "By signing below, both parties agree to the terms set forth in this Addendum, "
        "which is hereby incorporated into the original agreement referenced above.",
    )
    add_horizontal_rule(doc)

    add_signature_block(doc, "Client")
    add_horizontal_rule(doc)
    add_signature_block(doc, "Symphony Smart Homes")

    return doc


# ── CLI entry point ────────────────────────────────────────────────────────────

def parse_args():
    parser = argparse.ArgumentParser(
        description="Generate a Symphony Smart Homes Agreement Addendum (.docx)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    parser.add_argument("--client", required=True, help="Client full name")
    parser.add_argument("--project", required=True, help="Project name or address")
    parser.add_argument(
        "--items",
        required=True,
        help="Comma-separated list of client-supplied equipment items",
    )
    parser.add_argument(
        "--support-days",
        type=int,
        default=90,
        help="Post-installation defect coverage period in days (default: 90)",
    )
    parser.add_argument(
        "--integrations",
        required=True,
        help="Comma-separated list of third-party integrations (e.g., 'Control4 OS 3.x, Lutron RadioRA 3')",
    )
    parser.add_argument(
        "--output-dir",
        default=".",
        help="Directory to save the generated .docx file (default: current directory)",
    )
    return parser.parse_args()


def safe_filename(text):
    """Convert a string to a safe filename component."""
    return "".join(c if c.isalnum() or c in "-_" else "_" for c in text).strip("_")


def main():
    args = parse_args()

    # Validate
    if args.support_days < 1:
        print("ERROR: --support-days must be a positive integer.", file=sys.stderr)
        sys.exit(1)

    output_dir = os.path.abspath(args.output_dir)
    os.makedirs(output_dir, exist_ok=True)

    filename = f"addendum_{safe_filename(args.client)}_{safe_filename(args.project)}.docx"
    output_path = os.path.join(output_dir, filename)

    try:
        doc = build_document(
            client=args.client,
            project=args.project,
            items=args.items,
            support_days=args.support_days,
            integrations=args.integrations,
        )
        doc.save(output_path)
    except Exception as exc:
        print(f"ERROR: Failed to generate document — {exc}", file=sys.stderr)
        sys.exit(1)

    # Output path to stdout so callers can capture it
    print(output_path)


if __name__ == "__main__":
    main()
